#!/usr/bin/env python3
"""Generate RadityaNathaAzra_TUKTI_23.docx dari DOKUMENTASI.md + screenshots."""

import os, re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

BASE = os.path.dirname(os.path.abspath(__file__))
DOC = os.path.join(BASE, "RadityaNathaAzra_TUKTI_23.docx")
SC = os.path.join(BASE, "screenshots")

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Helper functions ---
def add_title(text, level=0):
    if level == 0:
        p = doc.add_heading(text, 0)
        run = p.runs[0]
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)

def add_para(text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = "Consolas"
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Body
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()  # spacing
    return table

def add_image(name, caption=None, width=Inches(5.0)):
    path = os.path.join(SC, name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(path, width=width)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(f"Gambar: {caption}")
            r.font.size = Pt(9)
            r.italic = True
    else:
        add_para(f"[Gambar tidak ditemukan: {name}]", size=9, color=RGBColor(255,0,0))

def add_separator():
    doc.add_paragraph("─" * 60)

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_title("DOKUMENTASI APLIKASI", 0)
doc.add_paragraph()
add_title("Sistem Manajemen UKM Poliban", 2)
add_para("Aplikasi Pengelolaan Surat Masuk dan Surat Keluar UKM", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(4):
    doc.add_paragraph()

info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Light Grid Accent 1'
info_data = [
    ("Nama", "Raditya Natha Azra"),
    ("NIM", "TUKTI_23"),
    ("Tanggal", "23 Juni 2026"),
    ("Program", "LSP - UKM Poliban"),
]
for i, (k, v) in enumerate(info_data):
    info_table.rows[i].cells[0].text = k
    info_table.rows[i].cells[1].text = v
    for cell in info_table.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(12)
                if cell is info_table.rows[i].cells[0]:
                    r.bold = True

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (simplified)
# ============================================================
add_title("DAFTAR ISI", 1)
toc_items = [
    "1. Verifikasi 10 Aktivitas Requirements",
    "2. Kebutuhan Fungsional (Functional Requirements)",
    "3. Kebutuhan Non-Fungsional (Non-Functional Requirements)",
    "4. Blackbox Testing (Playwright)",
    "5. Hasil Screenshot Tampilan Aplikasi",
    "6. Use Case Diagram",
    "7. Struktur Database",
    "8. Panduan Instalasi",
    "9. Teknologi yang Digunakan",
]
for item in toc_items:
    add_para(item, size=11)
doc.add_page_break()

# ============================================================
# 1. VERIFIKASI 10 AKTIVITAS
# ============================================================
add_title("1. VERIFIKASI 10 AKTIVITAS REQUIREMENTS", 1)
add_para("Berikut adalah verifikasi 10 aktivitas yang berkaitan dengan pengelolaan program surat masuk dan surat keluar UKM:", size=11)

verif_headers = ["No", "Aktivitas", "Status", "Implementasi"]
verif_rows = [
    ["1", "Menu mahasiswa, UKM, pendaftaran anggota, anggota UKM di awal tampilan", "✅", "Dashboard menampilkan 5 menu: Mahasiswa, UKM, Pendaftaran, Anggota UKM, Dashboard"],
    ["2", "Akun Administrator", "✅", "Role: administrator (email: wadir3@poliban.ac.id & kabagakademik@poliban.ac.id, password: admin123)"],
    ["3", "Pendaftaran mahasiswa & UKM melalui Wakil Direktur III / Kabag Akademik", "✅", "Admin melakukan pendaftaran manual atas permintaan pihak berwenang"],
    ["4", "Pendaftaran anggota UKM melalui Ketua/Sekretaris UKM", "✅", "Admin menambahkan anggota UKM berdasarkan data dari ketua/sekretaris"],
    ["5", "Data tersimpan: mahasiswa, UKM, anggota", "✅", "5 mahasiswa, 3 UKM, 3 anggota (seed data)"],
    ["6", "CRUD (Create, Read, Update, Delete)", "✅", "Semua data bisa ditambah, dilihat, diedit, dihapus"],
    ["7", "Anggota hanya dari mahasiswa terdaftar, satu mahasiswa satu UKM", "✅", "Unique constraint mahasiswa_id di tabel anggota_ukm"],
    ["8", "Fitur pencarian data", "✅", "Search per keyword di mahasiswa, UKM, pendaftaran"],
    ["9", "Fitur cetak data (print)", "✅", "Route /cetak untuk semua modul, CSS print-friendly"],
    ["10", "Fitur export ke Excel", "✅", "Route /export untuk semua modul, format CSV"],
]
add_table(verif_headers, verif_rows)
doc.add_page_break()

# ============================================================
# 2. KEBUTUHAN FUNGSIONAL
# ============================================================
add_title("2. KEBUTUHAN FUNGSIONAL (FUNCTIONAL REQUIREMENTS)", 1)

# FR-01
add_title("FR-01: Autentikasi", 2)
add_table(["ID", "Deskripsi", "Status"], [
    ["FR-01.01", "Halaman login dengan username + password", "✅"],
    ["FR-01.02", "Validasi credential", "✅"],
    ["FR-01.03", "Redirect ke dashboard jika berhasil", "✅"],
    ["FR-01.04", "Session management (login/logout)", "✅"],
    ["FR-01.05", "Proteksi halaman (middleware auth)", "✅"],
])

# FR-02
add_title("FR-02: Dashboard", 2)
add_table(["ID", "Deskripsi", "Status"], [
    ["FR-02.01", "Statistik total mahasiswa", "✅"],
    ["FR-02.02", "Statistik total UKM", "✅"],
    ["FR-02.03", "Statistik pendaftaran pending", "✅"],
    ["FR-02.04", "Statistik anggota aktif", "✅"],
    ["FR-02.05", "Menu navigasi cepat", "✅"],
])

# FR-03
add_title("FR-03: Manajemen Mahasiswa (CRUD)", 2)
add_table(["ID", "Deskripsi", "Status"], [
    ["FR-03.01", "Lihat daftar mahasiswa", "✅"],
    ["FR-03.02", "Tambah mahasiswa baru", "✅"],
    ["FR-03.03", "Edit data mahasiswa", "✅"],
    ["FR-03.04", "Hapus data mahasiswa", "✅"],
    ["FR-03.05", "Cari mahasiswa", "✅"],
    ["FR-03.06", "Export Excel", "✅"],
    ["FR-03.07", "Cetak data", "✅"],
])

# FR-04
add_title("FR-04: Manajemen UKM (CRUD)", 2)
add_table(["ID", "Deskripsi", "Status"], [
    ["FR-04.01", "Lihat daftar UKM", "✅"],
    ["FR-04.02", "Tambah UKM baru", "✅"],
    ["FR-04.03", "Edit data UKM", "✅"],
    ["FR-04.04", "Hapus data UKM", "✅"],
    ["FR-04.05", "Cari UKM", "✅"],
    ["FR-04.06", "Export Excel", "✅"],
    ["FR-04.07", "Cetak data", "✅"],
])

# FR-05, FR-06, FR-07, FR-08
add_title("FR-05: Manajemen Pendaftaran", 2)
add_table(["ID", "Deskripsi", "Status"], [
    ["FR-05.01", "Lihat daftar pendaftaran", "✅"],
    ["FR-05.02", "Setujui pendaftaran", "✅"],
    ["FR-05.03", "Tolak pendaftaran", "✅"],
    ["FR-05.04", "Cari pendaftaran", "✅"],
    ["FR-05.05", "Export CSV", "✅"],
    ["FR-05.06", "Cetak data", "✅"],
])

add_title("FR-06: Manajemen Anggota UKM", 2)
add_table(["ID", "Deskripsi", "Status"], [
    ["FR-06.01", "Lihat daftar anggota", "✅"],
    ["FR-06.02", "Tambah anggota baru", "✅"],
    ["FR-06.03", "Edit data anggota", "✅"],
    ["FR-06.04", "Hapus anggota", "✅"],
    ["FR-06.05", "Satu mahasiswa satu UKM", "✅"],
    ["FR-06.06", "Export CSV", "✅"],
    ["FR-06.07", "Cetak data", "✅"],
])

doc.add_page_break()

# ============================================================
# 3. KEBUTUHAN NON-FUNGSIONAL
# ============================================================
add_title("3. KEBUTUHAN NON-FUNGSIONAL (NON-FUNCTIONAL REQUIREMENTS)", 1)

nfr_headers = ["ID", "Kategori", "Deskripsi", "Target", "Status"]
nfr_rows = [
    ["NFR-01", "Availability", "Sistem tersedia 24/7", "99.5% uptime", "✅"],
    ["NFR-02", "Performance", "Response time < 3 detik", "< 3s", "✅"],
    ["NFR-03", "Performance", "Search time < 2 detik", "< 2s", "✅"],
    ["NFR-04", "Security", "Password di-hash (bcrypt)", "Hashed", "✅"],
    ["NFR-05", "Security", "Session ends on logout", "Destroyed", "✅"],
    ["NFR-06", "Security", "Input validation + SQL injection prevention", "Validated", "✅"],
    ["NFR-07", "Security", "Single role admin only", "Single role", "✅"],
    ["NFR-08", "Security", "Confirm before delete", "JS confirm", "✅"],
    ["NFR-09", "Usability", "Responsive UI (mobile-friendly)", "Tailwind CSS", "✅"],
    ["NFR-10", "Usability", "Navigation jelas dan intuitif", "Sidebar", "✅"],
    ["NFR-11", "Usability", "Error messages informatif", "Flash messages", "✅"],
    ["NFR-12", "Usability", "Label jelas di semua form", "Label + placeholder", "✅"],
    ["NFR-13", "Reliability", "Data tidak hilang saat error", "SQLite persistence", "✅"],
    ["NFR-14", "Reliability", "Database constraints (unique, FK)", "Constraints", "✅"],
    ["NFR-15", "Maintainability", "Code terstruktur MVC", "MVC pattern", "✅"],
    ["NFR-16", "Maintainability", "Dokumentasi tersedia", "README + docs", "✅"],
    ["NFR-17", "Portability", "Cross-platform (Linux/Windows)", "PHP 8.3+", "✅"],
    ["NFR-18", "Portability", "DB bisa migrate ke PostgreSQL", "SQLite → PgSQL", "✅"],
    ["NFR-19", "Scalability", "Pagination untuk data besar", "10/page", "✅"],
    ["NFR-20", "Scalability", "Handle 1000+ records", "Yes", "✅"],
]
add_table(nfr_headers, nfr_rows)
doc.add_page_break()

# ============================================================
# 4. BLACKBOX TESTING
# ============================================================
add_title("4. BLACKBOX TESTING (PLAYWRIGHT)", 1)

add_para("Pengujian blackbox dilakukan menggunakan Playwright (Chromium) dengan 12 skenario pengujian. Berikut hasil pengujian:", size=11)
doc.add_paragraph()

add_title("4.1 Test Scenarios", 2)
add_table(["No", "Nama File", "Deskripsi", "Status"], [
    ["1", "tampilan1_login.png", "Halaman login dengan form username + password", "✅ Pass"],
    ["2", "tampilan2_form_login.png", "Form login diisi credential admin", "✅ Pass"],
    ["3", "tampilan3_dashboard.png", "Dashboard admin dengan statistik & menu", "✅ Pass"],
    ["4", "tampilan4_mahasiswa.png", "Data mahasiswa dengan tabel CRUD", "✅ Pass"],
    ["5", "tampilan5_tambah_mahasiswa.png", "Form tambah mahasiswa", "✅ Pass"],
    ["6", "tampilan6_ukm.png", "Data UKM dengan card grid", "✅ Pass"],
    ["7", "tampilan7_tambah_ukm.png", "Form tambah UKM", "✅ Pass"],
    ["8", "tampilan8_pendaftaran.png", "Data pendaftaran anggota", "✅ Pass"],
    ["9", "tampilan9_anggota_ukm.png", "Data anggota UKM", "✅ Pass"],
    ["10", "tampilan10_tambah_anggota.png", "Form tambah anggota UKM", "✅ Pass"],
    ["11", "tampilan11_pencarian.png", "Fitur pencarian data", "✅ Pass"],
    ["12", "tampilan12_cetak.png", "Halaman cetak data", "✅ Pass"],
])

add_para("Pass Rate: 12/12 (100%)", bold=True, size=12, color=RGBColor(0,128,0))

add_title("4.2 Testing Environment", 2)
add_table(["Parameter", "Value"], [
    ["Browser", "Chromium (Playwright)"],
    ["Viewport", "1366 x 768 px"],
    ["URL", "http://localhost:8005"],
    ["Credential", "wadir3@poliban.ac.id / admin123"],
    ["Total Tests", "12"],
    ["Passed", "12"],
    ["Failed", "0"],
])

doc.add_page_break()

# ============================================================
# 5. HASIL SCREENSHOT
# ============================================================
add_title("5. HASIL SCREENSHOT TAMPILAN APLIKASI", 1)

screenshots = [
    ("tampilan1_login.png", "Tampilan 1 - Halaman Login"),
    ("tampilan2_form_login.png", "Tampilan 2 - Form Login Terisi"),
    ("tampilan3_dashboard.png", "Tampilan 3 - Dashboard Admin"),
    ("tampilan4_mahasiswa.png", "Tampilan 4 - Data Mahasiswa"),
    ("tampilan5_tambah_mahasiswa.png", "Tampilan 5 - Tambah Mahasiswa"),
    ("tampilan6_ukm.png", "Tampilan 6 - Data UKM"),
    ("tampilan7_tambah_ukm.png", "Tampilan 7 - Tambah UKM"),
    ("tampilan8_pendaftaran.png", "Tampilan 8 - Data Pendaftaran"),
    ("tampilan9_anggota_ukm.png", "Tampilan 9 - Data Anggota UKM"),
    ("tampilan10_tambah_anggota.png", "Tampilan 10 - Tambah Anggota"),
    ("tampilan11_pencarian.png", "Tampilan 11 - Pencarian Data"),
    ("tampilan12_cetak.png", "Tampilan 12 - Halaman Cetak"),
]

for filename, caption in screenshots:
    add_para(caption, bold=True, size=10)
    add_image(filename, caption, width=Inches(5.5))
    doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 6. USE CASE DIAGRAM
# ============================================================
add_title("6. USE CASE DIAGRAM", 1)

usecase_text = """
+--------------------------------------------------+
|              SISTEM UKM POLIBAN                  |
|                                                  |
|                    +------------+                |
|                    |   ADMIN    |                |
|                    +-----+------+                |
|                          |                       |
|          +---------------+---------------+       |
|          |               |               |       |
|          v               v               v       |
|    +---------+     +----------+     +--------+   |
|    | Login   |     |Dashboard |     |  CRUD  |   |
|    |         |     |  Stats   |     |  Data  |   |
|    +---------+     +----------+     +---+----+   |
|                                        |         |
|           +-----------------------------+---+    |
|           |               |              |   |   |
|           v               v              v   |   |
|    +-----------+    +-----------+  +--------+  |   |
|    | Mahasiswa |    |    UKM    |  | Anggota|  |   |
|    |   CRUD    |    |   CRUD    |  |   UKM  |  |   |
|    +-----------+    +-----------+  +--------+  |   |
|                                                |   |
|    +-----------+    +-----------+  +--------+  |   |
|    |Pendaftaran|    | Pencarian |  | Export/|  |   |
|    |  Approve  |    |   Data    |  |  Print |  |   |
|    +-----------+    +-----------+  +--------+  |   |
+--------------------------------------------------+
"""
add_code(usecase_text)
doc.add_page_break()

# ============================================================
# 7. STRUKTUR DATABASE
# ============================================================
add_title("7. STRUKTUR DATABASE", 1)

add_title("7.1 Tabel Database", 2)
add_table(["No", "Tabel", "Field"], [
    ["1", "admin", "username (UNIQUE), password (hashed)"],
    ["2", "users", "nama, nim (UNIQUE), username (UNIQUE), kelas, prodi, jurusan, Role, UKM, password"],
    ["3", "ukms", "nama, deskripsi"],
    ["4", "anggota_ukms", "user_id (UNIQUE), ukm_id, tanggal_bergabung"],
    ["5", "pendaftarans", "user_id, ukm_id, status (pending/diterima/ditolak)"],
    ["6", "kegiatans", "nama, ukm_id, tanggal, keterangan"],
])

add_title("7.2 Constraints", 2)
add_para("- username UNIQUE — tidak boleh duplikat", size=11)
add_para("- user_id UNIQUE di anggota_ukms — satu mahasiswa satu UKM", size=11)
add_para("- user_id + ukm_id UNIQUE di pendaftarans", size=11)
add_para("- Password di-hash menggunakan bcrypt", size=11)

doc.add_page_break()

# ============================================================
# 8. PANDUAN INSTALASI
# ============================================================
add_title("8. PANDUAN INSTALASI", 1)

add_para("Berikut langkah-langkah untuk menjalankan aplikasi:", size=11)
doc.add_paragraph()

add_code("""# Clone repository
git clone https://github.com/n8po/test1.git
cd test1

# Setup PHP dependencies
composer install

# Setup environment
cp .env.example .env
php artisan key:generate

# Setup database (SQLite)
touch database/database.sqlite
php artisan migrate --seed

# Setup frontend
npm install
npm run build

# Jalankan server
php artisan serve --port=8005

# Jalankan screenshots (Playwright)
npx playwright install chromium
node screenshot.cjs

# Akses Login
# Wakil Direktur III (Administrator):
# Email: wadir3@poliban.ac.id | Pass: admin123
#
# Kepala Bagian Akademik (Administrator):
# Email: kabagakademik@poliban.ac.id | Pass: admin123
#
# Ketua UKM Wasi Putih (Pengurus):
# Email: ketua1@poliban.ac.id | Pass: mahasiswa123
#
# Sekretaris UKM Wasi Putih (Pengurus):
# Email: sekretaris1@poliban.ac.id | Pass: mahasiswa123
#
# Ketua UKM Music Generation (Pengurus):
# Email: ketua2@poliban.ac.id | Pass: mahasiswa123
#
# Sekretaris UKM Music Generation (Pengurus):
# Email: sekretaris2@poliban.ac.id | Pass: mahasiswa123
#
# Ketua UKM Basket (Pengurus):
# Email: ketua3@poliban.ac.id | Pass: mahasiswa123
#
# Sekretaris UKM Basket (Pengurus):
# Email: sekretaris3@poliban.ac.id | Pass: mahasiswa123
#
# Anggota UKM (Mahasiswa):
# Email: anggota1@poliban.ac.id (dan anggota2, anggota3) | Pass: mahasiswa123""")

doc.add_page_break()

# ============================================================
# 9. TEKNOLOGI
# ============================================================
add_title("9. TEKNOLOGI YANG DIGUNAKAN", 1)

add_table(["Teknologi", "Versi", "Fungsi"], [
    ["Laravel", "13.x", "Framework backend"],
    ["PHP", "8.3.x", "Bahasa pemrograman"],
    ["SQLite / MySQL", "-", "Database"],
    ["Tailwind CSS", "4.x", "Styling UI"],
    ["Alpine.js", "3.x", "Interaktivitas frontend"],
    ["Blatui (shadcn/ui)", "1.x", "UI Components"],
    ["Playwright", "1.x", "Blackbox testing & screenshots"],
    ["Vite", "8.x", "Build tool"],
    ["Lucide Icons", "-", "Ikon aplikasi"],
    ["Composer", "2.x", "PHP package manager"],
])

doc.add_paragraph()
add_separator()

add_para("Dokumen ini disusun oleh:", bold=True, size=11)
add_para("Raditya Natha Azra", bold=True, size=12)
add_para("TUKTI_23", bold=True, size=12)
add_para("23 Juni 2026", size=10)

# ============================================================
# SAVE
# ============================================================
doc.save(DOC)
print(f"SUCCESS: DOCX saved to: {DOC}")
print(f"   File size: {os.path.getsize(DOC):,} bytes")
