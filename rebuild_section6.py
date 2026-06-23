import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from lxml import etree

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
body = doc.element.body

# Step 1: Find the key paragraph elements
tampilan12_end = None  # para 97 (Gambar: Tampilan 12)
blatui_heading = None  # "6. BLATUI..."
usecase_heading = None # "7. USE CASE..."

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Gambar: Tampilan 12' in t:
        tampilan12_end = p._element
    elif 'BLATUI UI' in t and p.style.name == 'Heading 1':
        blatui_heading = p._element
    elif 'USE CASE' in t and p.style.name == 'Heading 1':
        usecase_heading = p._element

print(f"Tampilan12 end: {tampilan12_end is not None}")
print(f"BlatUI heading: {blatui_heading is not None}")
print(f"Use case heading: {usecase_heading is not None}")

# Step 2: Remove all BlatUI content + Tampilan 13/14 + stray content between Tampilan12 end and Use Case
# Find the position after Tampilan 12 in body children
body_children = list(body)
t12_pos = body_children.index(tampilan12_end) if tampilan12_end in body_children else -1
uc_pos = body_children.index(usecase_heading) if usecase_heading in body_children else -1

print(f"T12 body pos: {t12_pos}, UC body pos: {uc_pos}")

# Remove elements between t12_pos+1 and uc_pos-1
removed = []
elements_to_remove = list(body_children[t12_pos+1:uc_pos])
print(f"Elements to remove between T12 and UC: {len(elements_to_remove)}")

for el in elements_to_remove:
    body.remove(el)
    removed.append(el)

print(f"Removed {len(removed)} elements")

# Step 3: Insert Tampilan 13 & 14 right after Tampilan 12
after_t12 = body_children[t12_pos + 1]  # This is now the first element after T12 (after removal)

new_elements = []

def add_new_para(text, is_img=False):
    p = doc.add_paragraph()
    if text:
        p.text = text
    if is_img:
        r = p.add_run()
        r.add_picture(text, width=Inches(5.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_elements.append(p._element)
    return p

# Tampilan 13
add_new_para('')
add_new_para('Tampilan 13 - Daftar Kegiatan')
add_new_para('')
add_new_para(os.path.join(SS, 'tampilan13_kegiatan.png'), is_img=True)
add_new_para('')
add_new_para('Gambar: Tampilan 13 - Daftar Kegiatan')
add_new_para('')

# Tampilan 14
add_new_para('Tampilan 14 - Profil Pengguna')
add_new_para('')
add_new_para(os.path.join(SS, 'tampilan14_profile.png'), is_img=True)
add_new_para('')
add_new_para('Gambar: Tampilan 14 - Profil Pengguna')
add_new_para('')
add_new_para('')

# Remove from end of body and insert after T12
next_after_t12 = list(body)[t12_pos + 1] if len(list(body)) > t12_pos + 1 else None
for el in reversed(new_elements):
    body.remove(el)
    if next_after_t12:
        body.insert(list(body).index(next_after_t12), el)
    else:
        body.append(el)

print(f"Inserted Tampilan 13/14 ({len(new_elements)} elements)")

# Step 4: Insert BlatUI section before Use Case
blatui_elements = []

# Heading
h1 = doc.add_paragraph('6. BLATUI UI COMPONENT LIBRARY')
h1.style = doc.styles['Heading 1']
blatui_elements.append(h1._element)

explanations = [
    'BlatUI adalah library UI component untuk Laravel Blade yang terinspirasi dari shadcn/ui. ' +
    'Library ini mengadopsi pendekatan copy-paste, di mana setiap component dapat diinstal satu per ' +
    'satu melalui CLI tanpa dependensi JavaScript runtime yang mengikat.',
    
    'BlatUI dibangun di atas stack BLAT (Blade, Laravel, Alpine.js, Tailwind CSS v4) dan mematuhi ' +
    'standar aksesibilitas WCAG AA. Setiap component diinstal dengan perintah: ' +
    'php artisan blatui:add <nama_component>',
    
    'Komponen yang digunakan dalam proyek UKM Poliban ini meliputi:',
    
    '- Button: Tombol dengan berbagai varian (default, secondary, outline, destructive, ghost)',
    '- Card: Container dengan header, content, dan footer untuk menampilkan informasi',
    '- Dialog & Alert Dialog: Modal window untuk form input dan konfirmasi hapus',
    '- Sidebar: Navigasi samping yang collapsible dengan icon mode',
    '- Table: Tabel responsif untuk menampilkan data mahasiswa, UKM, dan anggota',
    '- Avatar: Foto profil pengguna dengan fallback inisial',
    '- Badge: Label status dengan semantic tones (success, warning, danger, info)',
    '- Input, Textarea, Select, Combobox: Form controls dengan styling konsisten',
    '- Field & Field Error: Form field wrapper dengan label, input, dan error message',
    '- Alert: Callout notification dengan berbagai variant',
    '- Pagination: Navigasi halaman untuk data dengan banyak record',
    '- Separator: Pemisah visual antar elemen',
    '- Flip Card: Card yang membalik saat hover untuk menampilkan informasi tambahan',
    '- Combobox: Autocomplete dropdown untuk pemilihan mahasiswa dan jabatan',
    '- Alert Dialog: Konfirmasi destruktif untuk penghapusan data',
    
    'BlatUI menggunakan Alpine.js untuk interaktivitas frontend (seperti show/hide dialog, ' +
    'combobox autocomplete) dan Tailwind CSS v4 untuk styling. Komponen diinstal langsung ke ' +
    'direktori resources/views/components/ui/ sehingga kode sepenuhnya dimiliki dan dapat ' +
    'dikustomisasi tanpa batasan framework.',
    
    'Untuk menginstal komponen BlatUI, jalankan perintah berikut di terminal:',
    '',
    'composer require anousss007/blatui',
    'php artisan blatui:add button',
    'php artisan blatui:add card',
    'php artisan blatui:add dialog',
    'php artisan blatui:add sidebar',
    'php artisan blatui:add table',
    'php artisan blatui:add combobox',
    'php artisan blatui:add flip-card',
    '',
    'Setelah diinstal, komponen dapat digunakan di Blade view dengan prefix x-ui::, ' +
    'contoh: <x-ui::button>Tambah</x-ui::button>',
]

for text in reversed(explanations):
    p = doc.add_paragraph(text)
    blatui_elements.append(p._element)

# Move all BlatUI elements before Use Case
for el in blatui_elements:
    body.remove(el)
    body.insert(list(body).index(usecase_heading), el)

print(f"Inserted BlatUI section ({len(blatui_elements)} elements) before Use Case")

# Save
doc.save(DOCX)
print("SAVED!")
