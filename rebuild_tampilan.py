import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from lxml import etree
import os

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
body = doc.element.body

# Fix titles at paras 103 and 107
doc.paragraphs[103].text = 'Tampilan 13 - Daftar Kegiatan'
doc.paragraphs[107].text = 'Tampilan 14 - Profil Pengguna'

# Fix captions
doc.paragraphs[100].text = 'Gambar: Tampilan 14 - Profil Pengguna'
doc.paragraphs[104].text = 'Gambar: Tampilan 13 - Daftar Kegiatan'

# Now fix image content:
# Para 101 [IMG] should be tampilan13_kegiatan.png (Daftar Kegiatan)
# Para 105 [IMG] should be tampilan14_profile.png (Profil Pengguna)
# Currently they're swapped based on earlier fix

# Image para at 101 (currently has tampilan14_profile)
img_p101 = doc.paragraphs[101]
# Remove drawings
for r in img_p101.runs:
    for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        r._element.remove(d)
# Add correct image
run1 = img_p101.add_run()
run1.add_picture(os.path.join(SS, 'tampilan13_kegiatan.png'), width=Inches(5.5))
img_p101.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Image para at 105 (currently has tampilan13_kegiatan)
img_p105 = doc.paragraphs[105]
for r in img_p105.runs:
    for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        r._element.remove(d)
run2 = img_p105.add_run()
run2.add_picture(os.path.join(SS, 'tampilan14_profile.png'), width=Inches(5.5))
img_p105.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Now reorder: we want:
#   Tampilan 13 - Daftar Kegiatan
#   [blank]
#   [IMG 13]
#   [blank]
#   Gambar: Tampilan 13 - Daftar Kegiatan
#   [blank]
#   Tampilan 14 - Profil Pengguna
#   [blank]
#   [IMG 14]
#   [blank]
#   Gambar: Tampilan 14 - Profil Pengguna
#   [blank]

# Current order:
# 100: caption14  -> should move to after IMG14
# 101: IMG (correct for 13) -> stays, is IMG13
# 102: blank
# 103: title14 -> should be title13, need to move to 98 position
# 104: caption13 -> should move after IMG13
# 105: IMG (correct for 14) -> stays, is IMG14
# 106: blank
# 107: title13 -> should be title14, need to move to after caption14

# Elements to reorder
paras = [doc.paragraphs[i]._element for i in range(100, 108)]
for el in paras:
    body.remove(el)

# New order:
# title13 (was at 107, text now fixed)
# blank (was at 102)
# img13 (was at 101)
# blank (was at 106)
# caption13 (was at 104)
# blank (new)
# title14 (was at 103, text now fixed)
# blank (was at 102) - duplicate
# img14 (was at 105)
# blank (was at 106) - duplicate
# caption14 (was at 100)

# Actually let me take a simpler approach: just insert each at the correct position

# Get reference element (para 99 - blank before current messy section)
ref_el = doc.paragraphs[99]._element

# Build proper structure
new_order = []

# Tampilan 13 block
title13 = doc.add_paragraph('Tampilan 13 - Daftar Kegiatan')
new_order.append(title13._element)

blank1 = doc.add_paragraph()
new_order.append(blank1._element)

img13 = doc.add_paragraph()
run13 = img13.add_run()
run13.add_picture(os.path.join(SS, 'tampilan13_kegiatan.png'), width=Inches(5.5))
img13.alignment = WD_ALIGN_PARAGRAPH.CENTER
new_order.append(img13._element)

blank2 = doc.add_paragraph()
new_order.append(blank2._element)

cap13 = doc.add_paragraph('Gambar: Tampilan 13 - Daftar Kegiatan')
new_order.append(cap13._element)

blank3 = doc.add_paragraph()
new_order.append(blank3._element)

# Tampilan 14 block
title14 = doc.add_paragraph('Tampilan 14 - Profil Pengguna')
new_order.append(title14._element)

blank4 = doc.add_paragraph()
new_order.append(blank4._element)

img14 = doc.add_paragraph()
run14 = img14.add_run()
run14.add_picture(os.path.join(SS, 'tampilan14_profile.png'), width=Inches(5.5))
img14.alignment = WD_ALIGN_PARAGRAPH.CENTER
new_order.append(img14._element)

blank5 = doc.add_paragraph()
new_order.append(blank5._element)

cap14 = doc.add_paragraph('Gambar: Tampilan 14 - Profil Pengguna')
new_order.append(cap14._element)

blank6 = doc.add_paragraph()
new_order.append(blank6._element)

# Now find the position right before the BlatUI content (para 108)
ref_idx = 108
while ref_idx < len(doc.paragraphs) and doc.paragraphs[ref_idx].text.strip() == '':
    ref_idx += 1

# Actually, find the first BlatUI paragraph
blatui_start = None
for i in range(108, len(doc.paragraphs)):
    if 'BlatUI' in doc.paragraphs[i].text:
        blatui_start = doc.paragraphs[i]._element
        break

# Move all new elements before BlatUI start
if blatui_start:
    for el in new_order:
        body.remove(el)
        body.insert(list(body).index(blatui_start), el)

print("Reconstructed Tampilan 13 and 14 blocks")

doc.save(DOCX)
print("Saved!")
