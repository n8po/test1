import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from copy import deepcopy
from lxml import etree
import os

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)

body = doc.element.body
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def find_para(text_substr, style_name=None):
    """Find paragraph index by text content and optional style."""
    for i, p in enumerate(doc.paragraphs):
        if text_substr in p.text:
            if style_name is None or p.style.name == style_name:
                return i, p
    return None, None

def move_element(el, before_el):
    """Move el right before before_el."""
    parent = before_el.getparent()
    parent.remove(el)
    parent.insert(list(parent).index(before_el), el)

# ─── Move BlatUI section (para index ~153) before Section 7 (USE CASE at para 100) ───
blat_idx, blat_para = find_para('BLATUI UI', style_name='Heading 1')
usecase_idx, usecase_para = find_para('USE CASE', style_name='Heading 1')
teknologi_idx, teknologi_para = find_para('TEKNOLOGI YANG DIGUNAKAN', style_name='Heading 1')

print(f"BlatUI heading at: {blat_idx}")
print(f"Use Case heading at: {usecase_idx}")
print(f"Teknologi heading at: {teknologi_idx}")

if blat_idx and usecase_idx and blat_idx > usecase_idx:
    # Move all paragraphs from BlatUI heading to just before Teknologi back to before Use Case
    
    # Find all paragraph elements between BlatUI heading and Teknologi heading (exclusive)
    blat_el = doc.paragraphs[blat_idx]._element
    teku_el = doc.paragraphs[teknologi_idx]._element
    
    # Collect all elements from BlatUI to the element before Teknologi
    elements_to_move = []
    # Walk through body children
    found_blat = False
    for child in list(body):
        if child is blat_el:
            found_blat = True
        if found_blat:
            if child is teku_el:
                break
            if child.tag == f'{{{ns["w"]}}}p' or child.tag.endswith('}p'):
                elements_to_move.append(child)
    
    print(f"Elements to move: {len(elements_to_move)}")
    
    # Remove them
    for el in elements_to_move:
        body.remove(el)
    
    # Insert before Use Case element
    uc_el = doc.paragraphs[usecase_idx]._element
    for el in reversed(elements_to_move):
        body.insert(list(body).index(uc_el), el)
    
    print("  Moved BlatUI section before Use Case")
    
    # Re-fix heading text for Use Case (shift numbering)
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            t = p.text.strip()
            if 'USE CASE' in t.upper():
                p.text = '7. USE CASE DIAGRAM'
            elif 'STRUKTUR' in t.upper():
                p.text = '8. STRUKTUR DATABASE'
            elif 'AKUN' in t.upper() or 'CREDENTIAL' in t.upper():
                p.text = '9. DAFTAR AKUN LOGIN (CREDENTIAL)'
            elif 'PANDUAN' in t.upper():
                p.text = '10. PANDUAN INSTALASI'
            elif 'BLATUI' in t.upper():
                p.text = '6. BLATUI UI COMPONENT LIBRARY'
            elif 'TEKNOLOGI' in t.upper():
                p.text = '11. TEKNOLOGI YANG DIGUNAKAN'

# ─── Swap Tampilan 13 and 14 order ───
# Find title paragraphs
t13_title = t14_title = None
t13_img = t14_img = None
t13_cap = t14_cap = None

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Tampilan 13 - Daftar Kegiatan':
        t13_title = (i, p)
    elif p.text.strip() == 'Tampilan 14 - Profil Pengguna':
        t14_title = (i, p)

# Find image paragraphs for 13 and 14
img_paras = []
for i, p in enumerate(doc.paragraphs):
    has = any(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') for r in p.runs)
    if has and len(img_paras) < 14:
        img_paras.append(i)

print(f"Image indices: {img_paras}")

if len(img_paras) >= 14:
    # Para 12 = img 13, Para 13 = img 14
    img13_idx = img_paras[12]
    img14_idx = img_paras[13]
    
    # Also find caption paragraphs
    cap13 = None
    cap14 = None
    for i in range(img13_idx-1, max(img13_idx-3, -1), -1):
        p = doc.paragraphs[i]
        if 'Gambar:' in p.text and '13' in p.text:
            cap13 = i
            break
    for i in range(img14_idx-1, max(img14_idx-3, -1), -1):
        p = doc.paragraphs[i]
        if 'Gambar:' in p.text and '14' in p.text:
            cap14 = i
            break
    
    print(f"Img13: {img13_idx}, Cap13: {cap13}")
    print(f"Img14: {img14_idx}, Cap14: {cap14}")
    
    # Swap by moving elements
    if t13_title and t14_title:
        # We'll just reorder: collect all elements for tampilan 13 and 14, swap their positions
        t13_title_el = t13_title[1]._element
        t14_title_el = t14_title[1]._element
        
        # Swap the elements in the body
        t13_pos = list(body).index(t13_title_el)
        t14_pos = list(body).index(t14_title_el)
        
        # Find blocks: title, blank, img, blank, caption for each
        # For Tampilan 13: title at t13_pos, caption follows after img
        # Actually, let's just find the actual block ranges
        
        # Simpler approach: just take the last image (tampilan14) which should be profile
        # and move it before tampilan13
        
        # Actually the simplest: copy the correct images to the right paras
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Just ensure the images are correct
        p13 = doc.paragraphs[img13_idx]
        p14 = doc.paragraphs[img14_idx]
        
        # They already have the right images from the fix script
        # Just make sure the title order is correct by swapping text
        if t13_title and t14_title:
            t13_title[1].text = 'Tampilan 13 - Profil Pengguna'
            t14_title[1].text = 'Tampilan 14 - Daftar Kegiatan'
            
        # And swap captions
        if cap13 and cap14:
            # Swap image filenames in captions
            doc.paragraphs[cap13].text = 'Gambar: Tampilan 13 - Profil Pengguna'
            doc.paragraphs[cap14].text = 'Gambar: Tampilan 14 - Daftar Kegiatan'
        
        # Swap the actual image files
        from docx.shared import Inches
        for r in p13.runs:
            for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                r._element.remove(d)
        for r in p14.runs:
            for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                r._element.remove(d)
        
        run13 = p13.add_run()
        run13.add_picture(os.path.join(SS, 'tampilan14_profile.png'), width=Inches(5.5))
        p13.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run14 = p14.add_run()
        run14.add_picture(os.path.join(SS, 'tampilan13_kegiatan.png'), width=Inches(5.5))
        p14.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        print("  Swapped Tampilan 13 and 14 titles + images")

# Save
doc.save(DOCX)
print("\n✅ All order fixes applied!")
