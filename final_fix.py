import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)
body = doc.element.body
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# Fix 1: Swap the order of Tampilan 13 and 14 blocks (paras 100-110)
# Currently: caption14, img14, blank, title14, cap13, img13, blank, title13
# Target:     title13, blank, img13, blank, cap13, blank, title14, blank, img14, blank, cap14, blank

def has_drawing(p):
    return bool(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))

# Collect elements for tampilan 13 and 14
# Find the start (after Tampilan12 caption at para 97) and end (before BlatUI heading at para 112)
t12_cap_idx = None
blat_h_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'Gambar: Tampilan 12' in t:
        t12_cap_idx = i
    elif p.style.name == 'Heading 1' and 'BLATUI' in t:
        blat_h_idx = i
        break

print(f"T12 caption at: {t12_cap_idx}, BlatUI heading at: {blat_h_idx}")

# Collect all elements to replace
if t12_cap_idx and blat_h_idx:
    start_el = doc.paragraphs[t12_cap_idx + 1]._element  # first blank after T12 caption
    end_el = doc.paragraphs[blat_h_idx - 1]._element       # last para before BlatUI heading
    
    # Find positions in body
    body_children = list(body)
    start_pos = body_children.index(start_el)
    end_pos = body_children.index(end_el)
    
    print(f"Start body pos: {start_pos}, End body pos: {end_pos}")
    
    # Remove old elements
    old_elements = list(body_children[start_pos:end_pos+1])
    for el in old_elements:
        body.remove(el)
    
    print(f"Removed {len(old_elements)} old elements")
    
    # Build new content
    new_elements = []
    ref_el = None
    # Find what's now at start_pos
    if len(list(body)) > start_pos:
        ref_el = list(body)[start_pos]
    elif len(list(body)) > 0:
        ref_el = list(body)[-1]
    
    def add(text, is_img_path=None):
        p = doc.add_paragraph()
        if text:
            p.text = text
        if is_img_path:
            r = p.add_run()
            r.add_picture(is_img_path, width=Inches(5.5))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_elements.append(p._element)
        return p
    
    # Tampilan 13
    add('')
    add('Tampilan 13 - Daftar Kegiatan')
    add('')
    add('', os.path.join(SS, 'tampilan13_kegiatan.png'))
    add('')
    add('Gambar: Tampilan 13 - Daftar Kegiatan')
    add('')
    add('')
    
    # Tampilan 14
    add('Tampilan 14 - Profil Pengguna')
    add('')
    add('', os.path.join(SS, 'tampilan14_profile.png'))
    add('')
    add('Gambar: Tampilan 14 - Profil Pengguna')
    add('')
    
    # Insert before ref_el (or append if None)
    for el in reversed(new_elements):
        body.remove(el)
        if ref_el:
            body.insert(list(body).index(ref_el), el)
        else:
            body.append(el)
    
    print(f"Inserted {len(new_elements)} new elements")

# Fix 2: Fix BlatUI section content order
# Find BlatUI heading again after structure changed
blat_h_idx2 = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'BLATUI' in p.text:
        blat_h_idx2 = i
        break

# Find the USE CASE heading
uc_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'USE CASE' in p.text:
        uc_idx = i
        break

print(f"BlatUI heading at: {blat_h_idx2}, Use Case at: {uc_idx}")

if blat_h_idx2 and uc_idx:
    # Rebuild BlatUI content (currently reversed)
    # Remove all between BlatUI heading (inclusive) and Use Case heading (exclusive)
    start_el = doc.paragraphs[blat_h_idx2]._element
    end_el = doc.paragraphs[uc_idx]._element
    
    body_children = list(body)
    start_pos = body_children.index(start_el)
    end_pos = body_children.index(end_el)
    
    old_elements = list(body_children[start_pos:end_pos])
    for el in old_elements:
        body.remove(el)
    
    print(f"Removed {len(old_elements)} old BlatUI elements")
    
    # Build new BlatUI content in correct order
    new_el = []
    ref_el = list(body)[start_pos] if len(list(body)) > start_pos else None
    
    def badd(text):
        p = doc.add_paragraph(text)
        new_el.append(p._element)
    
    # Heading
    h1 = doc.add_paragraph('6. BLATUI UI COMPONENT LIBRARY')
    h1.style = doc.styles['Heading 1']
    new_el.append(h1._element)
    
    texts = [
        "BlatUI adalah library UI component untuk Laravel Blade yang terinspirasi dari shadcn/ui. Library ini mengadopsi pendekatan copy-paste, di mana setiap component dapat diinstal satu per satu melalui CLI tanpa dependensi JavaScript runtime yang mengikat.",
        "BlatUI dibangun di atas stack BLAT (Blade, Laravel, Alpine.js, Tailwind CSS v4) dan mematuhi standar aksesibilitas WCAG AA. Setiap component diinstal dengan perintah: php artisan blatui:add <nama_component>",
        "Komponen yang digunakan dalam proyek UKM Poliban ini meliputi:",
        "",
        "- Button: Tombol dengan berbagai varian (default, secondary, outline, destructive, ghost)",
        "- Card: Container dengan header, content, dan footer untuk menampilkan informasi",
        "- Dialog & Alert Dialog: Modal window untuk form input dan konfirmasi hapus",
        "- Sidebar: Navigasi samping yang collapsible dengan icon mode",
        "- Table: Tabel responsif untuk menampilkan data mahasiswa, UKM, dan anggota",
        "- Avatar: Foto profil pengguna dengan fallback inisial",
        "- Badge: Label status dengan semantic tones (success, warning, danger, info)",
        "- Input, Textarea, Select, Combobox: Form controls dengan styling konsisten",
        "- Field & Field Error: Form field wrapper dengan label, input, dan error message",
        "- Alert: Callout notification dengan berbagai variant",
        "- Pagination: Navigasi halaman untuk data dengan banyak record",
        "- Separator: Pemisah visual antar elemen",
        "- Flip Card: Card yang membalik saat hover untuk menampilkan informasi tambahan",
        "- Combobox: Autocomplete dropdown untuk pemilihan mahasiswa dan jabatan",
        "- Alert Dialog: Konfirmasi destruktif untuk penghapusan data",
        "",
        "BlatUI menggunakan Alpine.js untuk interaktivitas frontend (seperti show/hide dialog, combobox autocomplete) dan Tailwind CSS v4 untuk styling. Komponen diinstal langsung ke direktori resources/views/components/ui/ sehingga kode sepenuhnya dimiliki dan dapat dikustomisasi tanpa batasan framework.",
        "",
        "Untuk menginstal komponen BlatUI, jalankan perintah berikut di terminal:",
        "",
        "composer require anousss007/blatui",
        "php artisan blatui:add button",
        "php artisan blatui:add card",
        "php artisan blatui:add dialog",
        "php artisan blatui:add sidebar",
        "php artisan blatui:add table",
        "php artisan blatui:add combobox",
        "php artisan blatui:add flip-card",
        "",
        "Setelah diinstal, komponen dapat digunakan di Blade view dengan prefix x-ui::, contoh: <x-ui::button>Tambah</x-ui::button>",
    ]
    
    for t in texts:
        badd(t)
    
    # Insert before ref_el
    for el in reversed(new_el):
        body.remove(el)
        if ref_el:
            body.insert(list(body).index(ref_el), el)
        else:
            body.append(el)
    
    print(f"Inserted {len(new_el)} new BlatUI elements")

doc.save(DOCX)
print("\nFINAL FIX COMPLETE!")
