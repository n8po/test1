import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)

def has_drawing(p):
    for r in p.runs:
        if r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            return True
    return False

def clear_drawings(p):
    for r in p.runs:
        for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            r._element.remove(d)
        for o in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
            r._element.remove(o)

# ─── Step 1: Find all image paragraphs ───
img_idxs = []
for i, p in enumerate(doc.paragraphs):
    if has_drawing(p):
        img_idxs.append(i)
print(f"Image paragraphs: {img_idxs} ({len(img_idxs)} total)")

# ─── Step 2: Replace 12 existing images ───
new_images = [
    ('tampilan1_login.png', 'tampilan2_form_login.png', 'tampilan3_dashboard.png',
     'tampilan4_mahasiswa.png', 'tampilan5_tambah_mahasiswa.png', 'tampilan6_ukm.png',
     'tampilan7_tambah_ukm.png', 'tampilan8_pendaftaran.png', 'tampilan9_anggota_ukm.png',
     'tampilan10_tambah_anggota.png', 'tampilan11_pencarian.png', 'tampilan12_cetak.png')
]

for idx, para_idx in enumerate(img_idxs[:12]):
    p = doc.paragraphs[para_idx]
    clear_drawings(p)
    img_file = new_images[0][idx]
    img_path = os.path.join(SS, img_file)
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Update caption text beside it
        for j in range(para_idx-1, max(para_idx-3, -1), -1):
            cp = doc.paragraphs[j]
            if 'Gambar: Tampilan' in cp.text:
                # Extract number from filename
                num = img_file.replace('tampilan', '').replace('.png', '').split('_')[0]
                label = img_file.replace('tampilan', '').replace('.png', '')
                cp.text = f"Gambar: Tampilan {label} - {cp.text.split('-')[-1].strip() if '-' in cp.text else ''}"
                break
        print(f"  Replaced para {para_idx} with {img_file}")

# Tampilan 13 & 14: already inserted in previous run but order is swapped.
# Let's find their image paragraphs and make sure correct images are placed
# Get all image paragraphs again after modifications
img_idxs2 = []
for i, p in enumerate(doc.paragraphs):
    if has_drawing(p):
        img_idxs2.append(i)
print(f"Image paragraphs after replace: {img_idxs2}")

# Images 13 & 14 are the last two. Map them correctly
if len(img_idxs2) >= 14:
    # Image para for index 12 (0-based) = tampilan13
    p13 = doc.paragraphs[img_idxs2[12]]
    clear_drawings(p13)
    run13 = p13.add_run()
    run13.add_picture(os.path.join(SS, 'tampilan13_kegiatan.png'), width=Inches(5.5))
    p13.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f"  Fixed img para {img_idxs2[12]} -> tampilan13_kegiatan.png")
    
    # Image para for index 13 = tampilan14
    p14 = doc.paragraphs[img_idxs2[13]]
    clear_drawings(p14)
    run14 = p14.add_run()
    run14.add_picture(os.path.join(SS, 'tampilan14_profile.png'), width=Inches(5.5))
    p14.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f"  Fixed img para {img_idxs2[13]} -> tampilan14_profile.png")

# ─── Step 3: Fix section heading numbering ───
for p in doc.paragraphs:
    if p.style.name == 'Heading 1':
        t = p.text.strip()
        if t.startswith('5.'):
            p.text = '5. HASIL SCREENSHOT TAMPILAN APLIKASI (14 TAMPILAN)'
        elif 'USE CASE' in t.upper():
            p.text = '7. USE CASE DIAGRAM'
        elif 'STRUKTUR' in t.upper():
            p.text = '8. STRUKTUR DATABASE'
        elif 'AKUN' in t.upper() or 'CREDENTIAL' in t.upper():
            p.text = '9. DAFTAR AKUN LOGIN (CREDENTIAL)'
        elif 'PANDUAN' in t.upper():
            p.text = '10. PANDUAN INSTALASI'
        elif 'BLATUI' in t.upper():
            p.text = '6. BLATUI UI COMPONENT LIBRARY'
        elif 'TEKNOLOGI' in t.upper():
            p.text = '11. TEKNOLOGI YANG DIGUNAKAN'
    elif p.style.name == 'Heading 2':
        t = p.text.strip()
        if t.startswith('7.'):
            # These are under Struktur Database (section 8)
            p.text = '8.' + t[2:]

# ─── Step 4: Fix Daftar Isi ───
for p in doc.paragraphs:
    t = p.text.strip()
    if '5. Hasil Screenshot' in t:
        p.text = '5. Hasil Screenshot Tampilan Aplikasi (14 Tampilan)'
    elif t == '6. Use Case Diagram':
        p.text = '6. BlatUI UI Component Library'
    elif t == '7. Struktur Database':
        p.text = '7. Use Case Diagram'
    elif t == '8. Daftar Akun Login (Credential)':
        p.text = '8. Struktur Database'
    elif t == '9. Panduan Instalasi':
        p.text = '9. Daftar Akun Login (Credential)'
    elif t == '10. Teknologi yang Digunakan':
        p.text = '10. Panduan Instalasi'

# ─── Step 5: Update testing table ───
for table in doc.tables:
    h = ' '.join(c.text[:20] for c in table.rows[0].cells)
    if 'Nama File' in h and 'Deskripsi' in h:
        # Remove last 2 rows if they already exist, re-add
        while len(table.rows) > 13:
            tbl = table._tbl
            tbl.remove(table.rows[-1]._tr)
        for row_data in [
            ('13', 'tampilan13_kegiatan.png', 'Daftar kegiatan UKM dengan CRUD', '✅ Pass'),
            ('14', 'tampilan14_profile.png', 'Dialog profil pengguna dengan edit form', '✅ Pass'),
        ]:
            row = table.add_row()
            for ci, val in enumerate(row_data):
                row.cells[ci].text = val
        print("  Fixed testing table")

# ─── Step 6: Update total tests text ───
for p in doc.paragraphs:
    if '12/12' in p.text:
        p.text = p.text.replace('12/12', '14/14')

doc.save(DOCX)
print("\n✅ All fixes applied!")
