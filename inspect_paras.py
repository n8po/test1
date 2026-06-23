import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from lxml import etree
import os

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
body = doc.element.body

for i in range(98, 115):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    has_draw = bool(p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
    t = p.text.strip()
    flag = " [IMG]" if has_draw else "      "
    print(f'[{i}]{flag} {t[:100]}')
