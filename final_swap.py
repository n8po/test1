import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)
body = doc.element.body
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def remove_drawings(p):
    for r in p.runs:
        for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            r._element.remove(d)
        for o in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
            r._element.remove(o)

def set_image(p, filepath):
    remove_drawings(p)
    r = p.add_run()
    r.add_picture(filepath, width=Inches(5.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════════
# FIX 1: Swap Tampilan 13 and 14 labels + images
# Current structure (paras 99-110):
#   99: Gambar: Tampilan 14 - Profil Pengguna  → Gambar: Tampilan 13 - Daftar Kegiatan
#  101: [IMG] (currently tampilan14) → replace with tampilan13
#  103: Tampilan 14 - Profil Pengguna → Tampilan 13 - Daftar Kegiatan
#  106: Gambar: Tampilan 13 - Daftar Kegiatan → Gambar: Tampilan 14 - Profil Pengguna
#  108: [IMG] (currently tampilan13) → replace with tampilan14
#  110: Tampilan 13 - Daftar Kegiatan → Tampilan 14 - Profil Pengguna
# ══════════════════════════════════════════════════

# Fix texts
doc.paragraphs[99].text = 'Gambar: Tampilan 13 - Daftar Kegiatan'
doc.paragraphs[103].text = 'Tampilan 13 - Daftar Kegiatan'
doc.paragraphs[106].text = 'Gambar: Tampilan 14 - Profil Pengguna'
doc.paragraphs[110].text = 'Tampilan 14 - Profil Pengguna'

# Fix images
set_image(doc.paragraphs[101], os.path.join(SS, 'tampilan13_kegiatan.png'))
set_image(doc.paragraphs[108], os.path.join(SS, 'tampilan14_profile.png'))

print("Fixed Tampilan 13/14 labels and images")

# ══════════════════════════════════════════════════
# FIX 2: BlatUI content is BEFORE the heading
# Paras 112-145 = content (reversed), Para 146 = heading
# Move heading before content
# ══════════════════════════════════════════════════

blat_heading_el = doc.paragraphs[146]._element
blat_first_content = doc.paragraphs[112]._element

# Find positions in body
children = list(body)
h_pos = children.index(blat_heading_el)
fc_pos = children.index(blat_first_content)

print(f"BlatUI content at body pos {fc_pos}, heading at {h_pos}")

# Move heading element right before first content element
body.remove(blat_heading_el)
body.insert(list(body).index(blat_first_content), blat_heading_el)

print("Moved BlatUI heading before content")

# ══════════════════════════════════════════════════
# FIX 3: Reorder BlatUI content (currently reversed)
# The content paras 112-146 need to be reversed in order
# ══════════════════════════════════════════════════

# Find new positions after moving
blat_heading_idx = None
blat_end_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'BLATUI' in p.text:
        blat_heading_idx = i
    if p.style.name == 'Heading 1' and 'USE CASE' in p.text:
        blat_end_idx = i
        break

print(f"BlatUI heading now at: {blat_heading_idx}, Use Case at: {blat_end_idx}")

# Collect content between heading (exclusive) and Use Case (exclusive)
if blat_heading_idx and blat_end_idx and blat_end_idx > blat_heading_idx + 1:
    content_elements = []
    for i in range(blat_heading_idx + 1, blat_end_idx):
        el = doc.paragraphs[i]._element
        content_elements.append(el)
    
    # Remove all from body
    for el in content_elements:
        body.remove(el)
    
    # Re-insert in reverse order (after heading)
    ref_el = doc.paragraphs[blat_heading_idx]._element
    ref_pos = list(body).index(ref_el)
    
    for el in reversed(content_elements):
        body.insert(list(body).index(ref_el) + 1, el)
    
    print(f"Reversed {len(content_elements)} BlatUI content elements")

doc.save(DOCX)
print("FINAL SAVE DONE!")
