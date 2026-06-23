import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
import os

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
SS = r'C:\Project\latihan\lsp_ukm\test1\screenshots'

doc = Document(DOCX)

def has_drawing(p):
    for r in p.runs:
        if r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            return True
    return False

def clear_drawings(p):
    for r in p.runs:
        for d in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            r._element.remove(d)
        for o in r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}object'):
            r._element.remove(o)

# ─── Step 1: Find all image paragraphs by index ───
img_idxs = []
for i, p in enumerate(doc.paragraphs):
    if has_drawing(p):
        img_idxs.append(i)

print(f"Image paragraphs at indices: {img_idxs}")

# The 12 images should be at paragraphs 53, 57, 61, 65, 69, 73, 77, 81, 85, 89, 93, 97
# (even indices from the earlier listing)

# ─── Step 2: Replace images ───
new_images = [
    'tampilan1_login.png', 'tampilan2_form_login.png', 'tampilan3_dashboard.png',
    'tampilan4_mahasiswa.png', 'tampilan5_tambah_mahasiswa.png', 'tampilan6_ukm.png',
    'tampilan7_tambah_ukm.png', 'tampilan8_pendaftaran.png', 'tampilan9_anggota_ukm.png',
    'tampilan10_tambah_anggota.png', 'tampilan11_pencarian.png', 'tampilan12_cetak.png',
]

for idx, para_idx in enumerate(img_idxs[:12]):
    p = doc.paragraphs[para_idx]
    clear_drawings(p)
    img_path = os.path.join(SS, new_images[idx])
    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        print(f"  Replaced para {para_idx} with {new_images[idx]}")
    else:
        print(f"  WARNING: {img_path} not found!")

# ─── Step 3: Find Section 6 heading ───
sec6_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'USE CASE' in p.text.upper():
        sec6_idx = i
        break
print(f"Section 6 at index: {sec6_idx}")

# ─── Step 4: Insert Tampilan 13 & 14 before Section 6 ───
if sec6_idx:
    body = doc.paragraphs[sec6_idx]._element.getparent()
    sec6_el = doc.paragraphs[sec6_idx]._element
    
    new_entries = [
        ("Tampilan 13 - Daftar Kegiatan", "tampilan13_kegiatan.png", "Gambar: Tampilan 13 - Daftar Kegiatan"),
        ("Tampilan 14 - Profil Pengguna", "tampilan14_profile.png", "Gambar: Tampilan 14 - Profil Pengguna"),
    ]
    
    for title, img_file, caption in reversed(new_entries):
        img_path = os.path.join(SS, img_file)
        
        # We'll create paragraphs using doc.add_paragraph and then move them
        paras_to_move = []
        
        # Caption paragraph
        cap_p = doc.add_paragraph(caption)
        paras_to_move.append(cap_p._element)
        
        # Image paragraph
        if os.path.exists(img_path):
            img_p = doc.add_paragraph()
            img_run = img_p.add_run()
            img_run.add_picture(img_path, width=Inches(5.5))
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paras_to_move.append(img_p._element)
        
        # Blank para
        blank = doc.add_paragraph()
        paras_to_move.append(blank._element)
        
        # Title paragraph
        title_p = doc.add_paragraph(title)
        paras_to_move.append(title_p._element)
        
        # Move all before section 6
        for el in paras_to_move:
            body.remove(el)
            body.insert(list(body).index(sec6_el), el)
        
        print(f"  Inserted: {title}")

# ─── Step 5: Add BlatUI section before Teknologi ───
tech_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'TEKNOLOGI' in p.text.upper():
        tech_idx = i
        break
print(f"Teknologi section at: {tech_idx}")

if tech_idx:
    body = doc.paragraphs[tech_idx]._element.getparent()
    tech_el = doc.paragraphs[tech_idx]._element
    
    # Add heading and paragraphs in reverse order
    items_to_insert = []
    
    # Heading
    h1_p = doc.add_paragraph('11. BLATUI UI COMPONENT LIBRARY')
    h1_p.style = doc.styles['Heading 1']
    items_to_insert.append(h1_p._element)
    
    explanations = [
        "BlatUI adalah library UI component untuk Laravel Blade yang terinspirasi dari shadcn/ui. "
        "Library ini mengadopsi pendekatan copy-paste, di mana setiap component dapat diinstal satu per "
        "satu melalui CLI tanpa dependensi JavaScript runtime yang mengikat.",
        
        "BlatUI dibangun di atas stack BLAT (Blade, Laravel, Alpine.js, Tailwind CSS v4) dan mematuhi "
        "standar aksesibilitas WCAG AA. Setiap component diinstal dengan perintah: "
        "php artisan blatui:add <nama_component>",
        
        "Komponen yang digunakan dalam proyek UKM Poliban ini meliputi:",
        
        "- Button: Tombol dengan berbagai varian (default, secondary, outline, destructive, ghost)",
        "- Card: Container dengan header, content, dan footer untuk menampilkan informasi",
        "- Dialog & Alert Dialog: Modal window untuk form input dan konfirmasi hapus",
        "- Sidebar: Navigasi samping yang collapsible dengan icon mode",
        "- Table: Tabel responsif untuk menampilkan data mahasiswa, UKM, dan anggota",
        "- Avatar: Foto profil pengguna dengan fallback inisial",
        "- Badge: Label status dengan semantic tones (success, warning, danger, info)",
        "- Input, Textarea, Select, Combobox: Form controls dengan styling konsisten",
        "- Field & Field Error: Form field wrapper dengan label, input, dan error message",
        "- Alert: Callout notification dengan berbagai variant",
        "- Pagination: Navigasi halaman untuk data dengan banyak record",
        "- Separator: Pemisah visual antar elemen",
        "- Flip Card: Card yang membalik saat hover untuk menampilkan informasi tambahan",
        "- Combobox: Autocomplete dropdown untuk pemilihan mahasiswa dan jabatan",
        "- Alert Dialog: Konfirmasi destruktif untuk penghapusan data",
        
        "BlatUI menggunakan Alpine.js untuk interaktivitas frontend (seperti show/hide dialog, "
        "combobox autocomplete) dan Tailwind CSS v4 untuk styling. Komponen diinstal langsung ke "
        "direktori resources/views/components/ui/ sehingga kode sepenuhnya dimiliki dan dapat "
        "dikustomisasi tanpa batasan framework.",
        
        "Untuk menginstal komponen BlatUI, jalankan perintah berikut di terminal:",
        "",
        "composer require anousss007/blatui",
        "php artisan blatui:add button",
        "php artisan blatui:add card",
        "php artisan blatui:add dialog",
        "php artisan blatui:add sidebar",
        "php artisan blatui:add table",
        "php artisan blatui:add combobox",
        "php artisan blatui:add flip-card",
        "",
        "Setelah diinstal, komponen dapat digunakan di Blade view dengan prefix x-ui::, "
        "contoh: <x-ui::button>Tambah</x-ui::button>"
    ]
    
    for text in reversed(explanations):
        p = doc.add_paragraph(text)
        items_to_insert.append(p._element)
    
    # Move all items before Teknologi heading
    for el in items_to_insert:
        body.remove(el)
        body.insert(list(body).index(tech_el), el)
    
    print("  Inserted: BlatUI section")

# ─── Step 6: Update Daftar Isi (paragraphs 9-18) ───
for p in doc.paragraphs:
    t = p.text.strip()
    if t == '5. Hasil Screenshot Tampilan Aplikasi':
        p.text = '5. Hasil Screenshot Tampilan Aplikasi (14 Tampilan)'
    elif t == '6. Use Case Diagram':
        p.text = '6. BlatUI UI Component Library'
    elif t == '7. Struktur Database':
        p.text = '7. Use Case Diagram'
    elif t == '8. Daftar Akun Login (Credential)':
        p.text = '8. Struktur Database'
    elif t == '9. Panduan Instalasi':
        p.text = '9. Daftar Akun Login (Credential)'
    elif t == '10. Teknologi yang Digunakan':
        p.text = '10. Panduan Instalasi'

# Update actual Section 6 heading
for p in doc.paragraphs:
    if p.style.name == 'Heading 1' and 'USE CASE' in p.text.upper():
        # Renumber it to 7 since BlatUI is now 6
        if '7.' not in p.text:
            p.text = '7. ' + p.text.split('. ', 1)[1] if '. ' in p.text else p.text
    # Renumber other headings
    if p.style.name == 'Heading 1' and 'STRUKTUR' in p.text.upper():
        p.text = '8. STRUKTUR DATABASE'
    elif p.style.name == 'Heading 1' and 'AKUN' in p.text.upper() or 'DAFTAR' in p.text.upper() and 'Heading 1' in str(p.style):
        pass  # Skip

# ─── Step 7: Update testing table ───
for table in doc.tables:
    header_text = ' '.join(c.text for c in table.rows[0].cells)
    if 'Nama File' in header_text and 'Deskripsi' in header_text:
        table.rows[0].cells[0].text = 'No'
        # Add rows
        for row_data in [
            ('13', 'tampilan13_kegiatan.png', 'Daftar kegiatan UKM dengan CRUD', '✅ Pass'),
            ('14', 'tampilan14_profile.png', 'Dialog profil pengguna dengan edit form', '✅ Pass'),
        ]:
            row = table.add_row()
            for ci, val in enumerate(row_data):
                row.cells[ci].text = val
        print("  Updated testing table")

# ─── Step 8: Update total tests ───
for p in doc.paragraphs:
    if '12/12' in p.text:
        p.text = p.text.replace('12/12', '14/14')
        print("  Updated total test count")

# Save
doc.save(DOCX)
print("\n✅ DOCX updated successfully!")
