import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOCX = r'C:\Project\latihan\lsp_ukm\test1\RadityaNathaAzra_TUKTI_23_updated.docx'
doc = Document(DOCX)
body = doc.element.body

# Find BlatUI heading and Use Case heading
blat_h_idx = None
uc_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'BLATUI' in p.text:
        blat_h_idx = i
    if p.style.name == 'Heading 1' and 'USE CASE' in p.text:
        uc_idx = i
        break

print(f"Blat heading: {blat_h_idx}, Use case: {uc_idx}")

# Collect all elements between heading (inclusive) and Use Case (exclusive)
heading_el = doc.paragraphs[blat_h_idx]._element
uc_el = doc.paragraphs[uc_idx]._element

children = list(body)
h_pos = children.index(heading_el)
uc_pos = children.index(uc_el)

# The content = elements between heading+1 and uc-1
content_elements = list(children[h_pos + 1:uc_pos])
print(f"Content elements to reorder: {len(content_elements)}")

# Remove all content elements
for el in content_elements:
    body.remove(el)

# Re-insert in correct order (the reversed was already correct once, now they're backwards)
# Current order in paras: last content first, first content last
# Target: first content first, last content last
# So insert after heading in original order
for el in reversed(content_elements):  # reverse to get correct order
    body.insert(list(body).index(heading_el) + 1, el)

print("Reversed BlatUI content order")

doc.save(DOCX)
print("SAVED!")
